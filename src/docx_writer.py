from __future__ import annotations

import io
import re
from collections import Counter
from dataclasses import dataclass
from typing import List, Literal, Optional, Sequence, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from src.layout import LayoutAnalysis


def _docx_alignment(name: str) -> int:
    if name == "center":
        return WD_ALIGN_PARAGRAPH.CENTER
    if name == "right":
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT

def _normalize_lines(text: str) -> List[str]:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Keep empty lines (paragraph breaks), but trim right side.
    return [ln.rstrip() for ln in text.split("\n")]


def _join_wrapped_lines(lines: List[str]) -> str:
    # Join lines into a paragraph while respecting hyphenation.
    out = ""
    for ln in lines:
        if not out:
            out = ln
            continue
        if out.endswith("-"):
            out = out[:-1] + ln
        else:
            out = out + " " + ln
    return out.strip()


def _majority_bool(values: List[bool]) -> bool:
    if not values:
        return False
    return sum(1 for v in values if v) >= (len(values) / 2)


def _majority_alignment(values: List[str]) -> str:
    if not values:
        return "left"
    c = Counter(values)
    return c.most_common(1)[0][0]


MarkdownBlockType = Literal["paragraph", "heading", "ulist", "olist", "code", "table"]


@dataclass(frozen=True)
class MarkdownInline:
    text: str
    bold: bool = False
    italic: bool = False
    code: bool = False


@dataclass(frozen=True)
class MarkdownBlock:
    type: MarkdownBlockType
    source_line_idxs: List[int]  # indices into normalized lines
    text: str = ""
    items: Optional[List[str]] = None
    table: Optional[List[List[str]]] = None
    level: int = 0


def _parse_inline_markdown(text: str) -> List[MarkdownInline]:
    """
    Minimal inline markdown parser:
    - **bold**, __bold__
    - *italic*, _italic_
    - ***bold italic***, ___bold italic___
    - `inline code`
    - [text](url) -> rendered as "text (url)"
    """

    def parse_segment(seg: str, bold: bool = False, italic: bool = False) -> List[MarkdownInline]:
        out: List[MarkdownInline] = []
        i = 0
        while i < len(seg):
            # Inline code
            if seg[i] == "`":
                j = seg.find("`", i + 1)
                if j != -1:
                    code_text = seg[i + 1 : j]
                    if code_text:
                        out.append(MarkdownInline(text=code_text, bold=bold, italic=italic, code=True))
                    i = j + 1
                    continue

            # Links: [text](url)
            if seg[i] == "[":
                close = seg.find("]", i + 1)
                if close != -1 and close + 1 < len(seg) and seg[close + 1] == "(":
                    end = seg.find(")", close + 2)
                    if end != -1:
                        label = seg[i + 1 : close]
                        url = seg[close + 2 : end]
                        rendered = label if not url else f"{label} ({url})"
                        if rendered:
                            out.append(MarkdownInline(text=rendered, bold=bold, italic=italic))
                        i = end + 1
                        continue

            # Strong/em markers (recursive).
            for marker, b, it in (
                ("***", True, True),
                ("___", True, True),
                ("**", True, False),
                ("__", True, False),
                ("*", False, True),
                ("_", False, True),
            ):
                if seg.startswith(marker, i):
                    j = seg.find(marker, i + len(marker))
                    if j != -1:
                        inner = seg[i + len(marker) : j]
                        out.extend(parse_segment(inner, bold=bold or b, italic=italic or it))
                        i = j + len(marker)
                        break
            else:
                # Plain text until next special char.
                nxt = len(seg)
                for ch in ("`", "[", "*", "_"):
                    k = seg.find(ch, i + 1)
                    if k != -1:
                        nxt = min(nxt, k)
                out.append(MarkdownInline(text=seg[i:nxt], bold=bold, italic=italic))
                i = nxt
                continue

            continue

        return [t for t in out if t.text]

    if not any(ch in text for ch in ("`", "[", "*", "_")):
        return [MarkdownInline(text=text)]
    return parse_segment(text)


def _parse_markdown_blocks(lines: List[str]) -> List[MarkdownBlock]:
    blocks: List[MarkdownBlock] = []
    i = 0
    n = len(lines)

    def is_blank(s: str) -> bool:
        return s.strip() == ""

    heading_re = re.compile(r"^(#{1,6})\s+(.*)$")
    ulist_re = re.compile(r"^\s*[-*+]\s+(.*)$")
    olist_re = re.compile(r"^\s*(\d+)\.\s+(.*)$")
    table_sep_re = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")

    def looks_like_table_row(s: str) -> bool:
        # At least 2 pipes and not obviously a markdown link-only line.
        return s.count("|") >= 2

    def split_table_row(s: str) -> List[str]:
        parts = [p.strip() for p in s.strip().strip("|").split("|")]
        # Drop empty columns created by OCR noise at ends.
        while parts and parts[0] == "":
            parts = parts[1:]
        while parts and parts[-1] == "":
            parts = parts[:-1]
        return parts

    while i < n:
        ln = lines[i]
        if is_blank(ln):
            i += 1
            continue

        # Markdown/pipe table block
        if looks_like_table_row(ln):
            start = i
            row_lines: List[str] = []
            idxs: List[int] = []
            while i < n and not is_blank(lines[i]) and looks_like_table_row(lines[i]):
                row_lines.append(lines[i])
                idxs.append(i)
                i += 1

            # If it's too short, treat it as a normal paragraph.
            if len(row_lines) >= 2:
                rows: List[List[str]] = []
                for k, raw in enumerate(row_lines):
                    # Skip markdown separator lines like | --- | --- |
                    if table_sep_re.match(raw.strip()):
                        continue
                    cells = split_table_row(raw)
                    if cells:
                        rows.append(cells)
                if rows and max(len(r) for r in rows) >= 2:
                    blocks.append(MarkdownBlock(type="table", source_line_idxs=idxs, table=rows))
                    continue

            # Fallback: not a real table.
            i = start

        # Fenced code block
        if ln.strip().startswith("```"):
            start = i
            i += 1
            code_lines: List[str] = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < n and lines[i].strip().startswith("```"):
                i += 1
            blocks.append(
                MarkdownBlock(
                    type="code",
                    source_line_idxs=list(range(start, i)),
                    text="\n".join(code_lines).rstrip("\n"),
                )
            )
            continue

        # Heading
        m = heading_re.match(ln.strip())
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            blocks.append(MarkdownBlock(type="heading", source_line_idxs=[i], text=text, level=level))
            i += 1
            continue

        # Lists
        m_ul = ulist_re.match(ln)
        m_ol = olist_re.match(ln)
        if m_ul or m_ol:
            list_type: MarkdownBlockType = "ulist" if m_ul else "olist"
            items: List[str] = []
            idxs: List[int] = []
            while i < n:
                ln2 = lines[i]
                if is_blank(ln2):
                    break
                if list_type == "ulist":
                    m2 = ulist_re.match(ln2)
                    if not m2:
                        break
                    items.append(m2.group(1).strip())
                    idxs.append(i)
                else:
                    m2 = olist_re.match(ln2)
                    if not m2:
                        break
                    items.append(m2.group(2).strip())
                    idxs.append(i)
                i += 1
            blocks.append(MarkdownBlock(type=list_type, source_line_idxs=idxs, items=items))
            continue

        # Paragraph
        para_lines: List[str] = []
        idxs: List[int] = []
        while i < n and not is_blank(lines[i]):
            if lines[i].strip().startswith("```"):
                break
            if heading_re.match(lines[i].strip()):
                break
            if ulist_re.match(lines[i]) or olist_re.match(lines[i]):
                break
            para_lines.append(lines[i].strip())
            idxs.append(i)
            i += 1
        blocks.append(MarkdownBlock(type="paragraph", source_line_idxs=idxs, text=_join_wrapped_lines(para_lines)))

    return blocks


def _match_text_lines_to_layout_lines(text_lines: List[str], layout: LayoutAnalysis) -> List[Optional[int]]:
    """
    Returns an index mapping from each non-empty text line to a layout line index, or None.

    Because LightOnOCR returns text only (no boxes), this is a best-effort sequential mapping.
    """
    layout_lines = layout.lines
    non_empty = [i for i, ln in enumerate(text_lines) if ln.strip() != ""]
    mapping: List[Optional[int]] = [None] * len(non_empty)

    # Ideal case: same number of lines.
    if len(non_empty) == len(layout_lines):
        for k, _ in enumerate(non_empty):
            mapping[k] = k
        return mapping

    # Fallback: map sequentially up to the smaller count.
    m = min(len(non_empty), len(layout_lines))
    for k in range(m):
        mapping[k] = k
    return mapping


def _line_index_to_layout_index(text_lines: List[str], layout: LayoutAnalysis) -> List[Optional[int]]:
    non_empty_positions = [i for i, ln in enumerate(text_lines) if ln.strip() != ""]
    non_empty_to_layout = _match_text_lines_to_layout_lines(text_lines, layout)
    out: List[Optional[int]] = [None] * len(text_lines)
    for k, pos in enumerate(non_empty_positions):
        if k < len(non_empty_to_layout):
            out[pos] = non_empty_to_layout[k]
    return out


def _styles_from_layout(
    line_idxs: Sequence[int],
    text_line_to_layout: Sequence[Optional[int]],
    layout: LayoutAnalysis,
) -> Tuple[str, bool, bool]:
    aligns: List[str] = []
    bolds: List[bool] = []
    italics: List[bool] = []
    for tli in line_idxs:
        if tli < 0 or tli >= len(text_line_to_layout):
            continue
        li = text_line_to_layout[tli]
        if li is None or li >= len(layout.lines):
            continue
        s = layout.lines[li].style
        aligns.append(s.alignment)
        bolds.append(bool(s.bold))
        italics.append(bool(s.italic))
    return _majority_alignment(aligns), _majority_bool(bolds), _majority_bool(italics)


def _add_inline_runs(paragraph, text: str, default_bold: bool, default_italic: bool) -> None:
    # Apply heuristic bold/italic only if the text doesn't already contain markdown emphasis markers.
    apply_defaults = default_bold or default_italic
    if re.search(r"(\*\*|__|\*|_)", text) or "`" in text:
        apply_defaults = False

    tokens = _parse_inline_markdown(text)
    for tok in tokens:
        run = paragraph.add_run(tok.text)
        if tok.bold:
            run.bold = True
        elif apply_defaults and default_bold:
            run.bold = True

        if tok.italic:
            run.italic = True
        elif apply_defaults and default_italic:
            run.italic = True

        if tok.code:
            run.font.name = "Consolas"
            run.font.size = Pt(10)


def build_docx(title: str, ocr_text: str, layout: LayoutAnalysis) -> bytes:
    doc = Document()

    if title.strip():
        doc.add_heading(title.strip(), level=1)

    text_lines = _normalize_lines(ocr_text)
    text_line_to_layout = _line_index_to_layout_index(text_lines, layout)
    blocks = _parse_markdown_blocks(text_lines)

    for block in blocks:
        align, default_bold, default_italic = _styles_from_layout(block.source_line_idxs, text_line_to_layout, layout)

        if block.type == "table" and block.table:
            rows = block.table
            cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                for c_idx in range(cols):
                    cell = table.cell(r_idx, c_idx)
                    cell.text = ""
                    para = cell.paragraphs[0]
                    para.alignment = _docx_alignment(align)
                    txt = row[c_idx] if c_idx < len(row) else ""
                    _add_inline_runs(para, txt, default_bold=default_bold, default_italic=default_italic)
            continue

        if block.type == "heading":
            level = max(1, min(6, int(block.level or 1)))
            para = doc.add_heading(block.text.strip(), level=level)
            para.alignment = _docx_alignment(align)
            continue

        if block.type in ("ulist", "olist") and block.items:
            style_name = "List Bullet" if block.type == "ulist" else "List Number"
            for item in block.items:
                para = doc.add_paragraph(style=style_name)
                para.alignment = _docx_alignment(align)
                _add_inline_runs(para, item, default_bold=default_bold, default_italic=default_italic)
            continue

        if block.type == "code":
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(block.text)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            continue

        para = doc.add_paragraph()
        para.alignment = _docx_alignment(align)
        _add_inline_runs(para, block.text, default_bold=default_bold, default_italic=default_italic)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

